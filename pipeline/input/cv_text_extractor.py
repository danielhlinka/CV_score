from __future__ import annotations

from pathlib import Path

import pdfplumber
from docx import Document

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def extract_text(file_path: str) -> str:
    """Extract plain text from supported CV file types.
    Dispatches to format-specific readers with explicit extension checks."""
    extension = Path(file_path).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {extension}. Use PDF or DOCX.")
    if extension == ".pdf":
        return _extract_from_pdf(file_path)
    return _extract_from_docx(file_path)


def _extract_from_pdf(file_path: str) -> str:
    """Read PDF pages and concatenate non-empty extracted text blocks.
    Raises when no text is available, e.g., scanned image-only files."""
    text_parts: list[str] = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    if not text_parts:
        raise ValueError("Could not extract any text from PDF. File may be scanned/image-based.")
    return "\n".join(text_parts)


def _extract_from_docx(file_path: str) -> str:
    """Read textual content from DOCX paragraphs and table cells.
    Returns a newline-separated representation for downstream parsing."""
    doc = Document(file_path)
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " — ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)
