import os
import re
import pdfplumber
from docx import Document

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

def extract_text(file_path: str) -> str:
    """Extract raw text from a PDF or DOCX file."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}. Use PDF or DOCX.")
    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext == ".docx":
        return _extract_from_docx(file_path)

def _extract_from_pdf(file_path: str) -> str:
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    if not text_parts:
        raise ValueError("Could not extract any text from PDF. File may be scanned/image-based.")
    return "\n".join(text_parts)

def _extract_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    parts = []

    # Read paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Read tables too ← THIS IS THE MISSING PIECE
    for table in doc.tables:
        for row in table.rows:
            row_text = " — ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)

def validate_text(text: str, min_chars: int = 200) -> bool:
    """Sanity check — is there enough text to work with?"""
    return len(text.strip()) >= min_chars


SECTION_HEADERS = {
    "work": r'work experience|experience|employment|pracovné skúsenosti|zamestnanie',
    "skills": r'skills|technical skills|zručnosti|znalosti',
    "education": r'education|vzdelanie|qualifications',
}


def extract_sections(text: str) -> dict:
    sections = {"work": "", "skills": "", "education": "", "other": ""}

    # find all header positions
    header_pattern = r'(' + '|'.join(SECTION_HEADERS.values()) + r')'
    splits = re.split(header_pattern, text, flags=re.IGNORECASE)

    current = "other"
    for chunk in splits:
        matched = False
        for section, pattern in SECTION_HEADERS.items():
            if re.fullmatch(pattern, chunk.strip(), re.IGNORECASE):
                current = section
                matched = True
                break
        if not matched:
            sections[current] += chunk

    return sections

def extract_education(text: str) -> str:
    text_lower = text.lower()

    if any(k in text_lower for k in ["phd", "ph.d", "doc.", "prof.", "doctorate"]):
        return "phd"
    if any(k in text_lower for k in ["master", "msc", "m.sc", "ing.", "mgr.", "mba"]):
        return "master"
    if any(k in text_lower for k in ["university", "univerzita", "bachelor", "bsc", "b.sc", "utb"]):
        return "bachelor"
    if any(k in text_lower for k in ["secondary", "high school", "spše", "stredná", "gymnázium"]):
        return "high_school"

    return "none"

SKILL_GROUPS = {
    # Management & leadership (grouped to avoid duplicates)
    "leadership":           ["leadership", "lead", "led"],
    "people management":    ["people management", "team management", "team lead", "managed"],
    "mentoring":            ["mentoring", "mentored"],
    "stakeholder management": ["stakeholder", "stakeholders"],
    "strategic planning":   ["strategic planning", "strategy"],
}

KNOWN_SKILLS = [
    # Technical
    "python", "java", "php", "html", "css", "javascript", "c", "c++", "c#",
    "sql", "postgresql", "mysql", "mongodb", "flask", "django", "react",
    "docker", "git", "linux", "blender", "3d modeling", "animation",
    "microsoft word", "excel", "powerpoint", "typescript", "node", "aws",
    # Management (ungrouped — no synonyms)
    "hiring", "budget", "p&l", "cross-functional", "roadmap", "okr",
    "product management", "project management", "risk management",
    # Delivery & process
    "agile", "scrum", "kanban", "sprint", "delivery",
]

def extract_skills(text: str) -> list:
    text_lower = text.lower()
    found = []
    seen = set()

    # Grouped skills — return canonical name if any variant matches
    for canonical, variants in SKILL_GROUPS.items():
        for variant in variants:
            if re.search(rf'\b{re.escape(variant)}\b', text_lower):
                found.append(canonical)
                seen.add(canonical)
                break  # first match wins, skip other variants

    # Ungrouped skills from KNOWN_SKILLS
    for skill in KNOWN_SKILLS:
        if re.search(rf'\b{re.escape(skill)}\b', text_lower):
            found.append(skill)

    return found